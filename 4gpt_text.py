from openai import OpenAI
from dotenv import load_dotenv
import os

load_dotenv()

summary_prompt = """당신은 언어 이해 및 요약에 훈련된 고도로 숙련된 AI입니다.
다음 텍스트를 읽고 간결한 추상적 단락으로 요약했으면 합니다.
전체 텍스트를 읽을 필요 없이 토론의 요점을 이해하는 데 도움이 될 수 있는 일관되고
읽을 수 있는 요약을 제공하여 가장 중요한 요점을 유지하는 것을 목표로 합니다.
불필요한 세부 사항이나 접선 사항은 피하십시오."""
key_points_prompt = """당신은 정보를 핵심 포인트로 전달하는 데 특화된 능숙한 AI입니다.
다음 텍스트를 기반으로 논의되거나 언급된 주요 포인트를 확인하고 나열합니다.
이는 논의의 본질에 가장 중요한 아이디어, 결과 또는 주제가 되어야 합니다.
당신의 목표는 누군가가 읽을 수 있는 목록을 제공하여 이야기된 내용을 빠르게 이해하는 것입니다."""
action_items_prompt = """당신은 대화를 분석하고 행동 항목을 추출하는 데 있어 AI 전문가입니다.
본문을 검토하고 합의되거나 수행이 필요하다고 언급된 모든 작업, 과제 또는 행동을 식별하십시오.
이것들은 특정 개인에게 할당된 작업일 수도 있고 그룹이 취하기로 결정한 일반적인 행동일 수도 있습니다.
이러한 행동 항목을 명확하고 간결하게 나열하십시오."""
sentiment_prompt = """당신은 언어와 감정 분석에 전문성을 갖춘 AI로서 당신의 과제는 다음 텍스트의 감
정을 분석하는 것입니다.
토론의 전체적인 톤, 사용된 언어가 전달하는 감정, 단어와 구가 사용되는 맥락을 고려하십시오.
감정이 일반적으로 긍정적인지 부정적인지 중립적인지를 표시하고 가능한 한 당신의 분석에 대해 간략한 설명을 제
공하십시오."""

os.environ.get("OPEN_API_KEY")


api_key = os.environ.get("OPEN_API_KEY")
client = OpenAI(
    api_key= api_key
)
def text_extraction(text, prompt):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
        { "role": "system", "content": prompt },
        { "role": "user", "content": text }
        ]
    )
    return response.choices[0].message.content

audio_file_path = "assets/news.mp3"
with open(audio_file_path, 'rb') as audio_file:
    transcription = client.audio.transcriptions.create(
    model="whisper-1",
    file=audio_file,
    response_format="text"
)
    
abstract_summary = text_extraction(transcription, summary_prompt)
key_points = text_extraction(transcription, key_points_prompt)
action_items = text_extraction(transcription, action_items_prompt)
sentiment = text_extraction(transcription, sentiment_prompt)
news_data = {
    'abstract_summary': abstract_summary,
    'key_points': key_points,
    'action_items': action_items,
    'sentiment': sentiment
}
print(news_data)

from docx import Document
doc = Document()

for key,values in news_data.items():
    heading = ' ',join(word.capitalize() for word in key.split('_'))
    doc.add_heading(heading, level=1)
    doc.add_paragraph(value)
    doc.add_paragraph()
doc.save('new_minutes.docx')